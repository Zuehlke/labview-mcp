#!/usr/bin/env python3
"""Generate a modern Word (.docx) documentation for a LabVIEW library, class or project.

Usage:
    py generate_labview_doc.py <data.json> <output.docx>
        [--structure-out <png>] [--uml-out <png>] [--browser <exe>]

Requirements: python-docx (required); Pillow (optional — used to read the pixel
size of the icon/connector-pane images so they are placed at their natural
aspect ratio; without it they are placed at fixed widths). A Chromium browser
(Edge/Chrome) must be installed for rendering the two diagrams (SVG -> PNG
headless).

The document:
    title + meta line + short description
    structure diagram   (library -> folders with access scope -> items)
    UML class diagram   (only when "classes" is non-empty)
    table of contents   (real Word TOC field — Word updates it on first open)
    one section per PUBLIC VI: icon + connector pane, description, terminal table
    appendix: non-public members, unreadable files, missing files

All styling lives here. The caller supplies DATA, never layout.

JSON contract (UTF-8):
{
  "title":     "Module1",                   # document title
  "language":  "en",                        # "de" or "en"; omitted defaults to "en"
  "generated": "2026-08-06",                # date shown in the meta line
  "labels":    { },                         # optional per-key label overrides
  "target": {
    "path":        "C:\\...\\SampleModule.lvlib",
    "kind":        "library",               # library | class | project | folder
    "description": "What this library is for.",
    "version":     "1.0.0.0",               # optional
    "locked":      false,                   # optional
    "labview":     "2026 (26008000)"        # optional
  },
  "structure": [                            # ordered tree, file order preserved
    { "name": "Public API", "kind": "folder", "scope": "public", "children": [
        { "name": "Start Module.vi", "kind": "vi", "scope": "public",
          "documented": true },
        { "name": "Module Data--cluster.ctl", "kind": "ctl", "scope": "public",
          "documented": false, "unreadable": "typedef" }
    ]},
    { "name": "Main.vi", "kind": "vi", "scope": "private", "documented": false }
  ],
  "classes": [                              # omit or leave empty for non-OOP
    { "name": "Derived.lvclass", "parent": "Base.lvclass",
      "parent_library": "", "external_parent": false,
      "private_data": "Derived.ctl",
      "methods": [ { "name": "Init.vi", "scope": "public",
                     "dynamic_dispatch": true } ] }
  ],
  "vis": [                                  # PUBLIC VIs only — one section each
    { "name": "Start Module.vi",
      "qualified_name": "SampleModule.lvlib:Start Module.vi",
      "path": "C:\\...\\Start Module.vi",
      "scope": "public",
      "description": "Starts the module and returns its Module ID.",
      "description_derived": false,
      "icon":    "C:\\temp\\images\\Start Module_icon.png",     # optional
      "conpane": "C:\\temp\\images\\Start Module_conpane.png",  # optional
      "terminals": [
        { "name": "error in", "type": "cluster", "conIdx": 3,
          "direction": "input", "default": "no error", "description": "" }
      ] }
  ],
  "non_public": [ { "name": "Main.vi", "scope": "private", "folder": "Private" } ],
  "unreadable": [ { "name": "X.ctl", "reason": "typedef" } ],
  "missing_files": [ "C:\\...\\Gone.vi" ],
  "notes": [ "LabVIEW was started by this run." ]
}

Kinds understood in "structure": folder, library, class, vi, vim, ctl, other.
Scopes understood everywhere: public, private, protected, community, unknown.
"""
import json
import math
import os
import shutil
import subprocess
import sys
import tempfile
from xml.sax.saxutils import escape as xml_escape

try:
    from PIL import Image
except ImportError:  # without Pillow the images are placed at fixed widths
    Image = None

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt, RGBColor

try:
    from docx.oxml import OxmlElement
except ImportError:  # python-docx >= 1.2 moved it
    from docx.oxml.parser import OxmlElement

FONT = "Segoe UI"

# Every table in the document is set at this size. Terminal tables are the widest
# thing here — six columns of type names and descriptions — so they set the floor.
TABLE_PT = 8

# Language used when the data JSON does not name one.
DEFAULT_LANGUAGE = "en"

# Restrained, modern palette — dark petrol accent, light fills, no icons.
# Deliberately identical to generate_teststand_doc.py so documents from both
# MCP servers look like one family.
C_INK = "333333"
C_MUTED = "6E7B84"
C_ACCENT = "1F4E5F"
C_ACCENT2 = "2E6E80"
C_RULE = "C9D6DC"
C_TBL_LINE = "D9E2E7"
C_TBL_LINE2 = "B9C6CD"
C_NODE_FILL = "F2F6F8"
C_NODE_STROKE = "9FB6C1"
C_EDGE = "5B7A8C"

# Access scope is the one thing this document is opinionated about: public code
# carries the accent, everything else recedes. No red/green — a private VI is
# not an error.
SCOPE_COLOR = {
    "public": C_ACCENT,
    "private": "97A6AE",
    "protected": "7E8FA8",
    "community": "8A8296",
    "unknown": "AAB4B9",
}
SCOPE_MARK = {"public": "+", "protected": "#", "private": "-", "community": "~", "unknown": "?"}

# Item kinds in the structure diagram: a small square in this color.
KIND_COLOR = {
    "vi": C_ACCENT2,
    "vim": "5E8F7E",
    "ctl": "9A8B5F",
    "class": "6A5F8F",
    "library": C_ACCENT,
    "other": "9FB6C1",
}

LABELS = {
    "de": {
        "toc": "Inhaltsverzeichnis",
        "toc_placeholder": "Das Inhaltsverzeichnis wird beim ersten Öffnen in Word aktualisiert (F9).",
        "structure": "Struktur",
        "uml": "Klassendiagramm",
        "public_vis": "Öffentliche VIs",
        "appendix": "Anhang",
        "no_description": "Keine Beschreibung hinterlegt.",
        "derived": "(abgeleitet)",
        "version": "Version",
        "locked": "gesperrt",
        "generated_on": "erstellt",
        "structure_intro": "Das folgende Diagramm zeigt den Aufbau von {name}: Ordner mit ihrem "
                           "Zugriffsbereich, die enthaltenen VIs, Typdefinitionen und Klassen.",
        "uml_intro": "Das folgende Klassendiagramm zeigt die Vererbungsbeziehungen und die "
                     "Methoden je Klasse.",
        "terminals": "Anschlüsse",
        "no_terminals": "Keine Anschlüsse am Connector Pane.",
        "t_name": "Name",
        "t_type": "Typ",
        "t_dir": "Richtung",
        "t_default": "Standardwert",
        "t_desc": "Beschreibung",
        "dir_in": "Eingang",
        "dir_out": "Ausgang",
        "fp_only": "nur Frontpanel",
        "req_required": "erforderlich",
        "req_recommended": "empfohlen",
        "req_optional": "optional",
        "non_public": "Nicht öffentliche Elemente",
        "non_public_intro": "Diese Elemente sind nicht öffentlich und werden daher nicht "
                            "einzeln dokumentiert.",
        "unreadable": "Nicht lesbare Dateien",
        "missing": "Fehlende Dateien",
        "notes": "Hinweise zum Lauf",
        "np_name": "Name",
        "np_scope": "Zugriff",
        "np_folder": "Ordner",
        "u_reason": "Grund",
        "reason_typedef": "Typdefinition (.ctl) — von LabVIEW nicht als VI lesbar",
        "reason_password": "passwortgeschützt",
        "reason_diagram": "Blockdiagramm nicht lesbar",
        "short_password": "geschützt",
        "short_diagram": "Diagramm n. lesbar",
        "scope_public": "Öffentlich",
        "scope_private": "Privat",
        "scope_protected": "Geschützt",
        "scope_community": "Community",
        "scope_unknown": "unbekannt",
        "legend_folder": "Ordner",
        "legend_vi": "VI",
        "legend_vim": "Malleable VI",
        "legend_ctl": "Typdefinition",
        "legend_class": "Klasse",
        "legend_inherits": "erbt von",
        "legend_external": "externe Klasse",
        "continued": "Fortsetzung",
        "collapsed": "{n} Elemente — siehe Anhang",
        "structure_collapsed_note": "Nicht öffentliche Ordner sind zusammengefasst dargestellt; "
                                    "ihre Elemente sind im Anhang einzeln aufgeführt.",
        "more": "{n} weitere",
        "page": "Seite",
        "page_of": "von",
        "items_count": "Elemente",
        "no_images": "Icon und Connector Pane konnten nicht erzeugt werden.",
    },
    "en": {
        "toc": "Table of Contents",
        "toc_placeholder": "The table of contents is updated when the document is first opened in Word (F9).",
        "structure": "Structure",
        "uml": "Class Diagram",
        "public_vis": "Public VIs",
        "appendix": "Appendix",
        "no_description": "No description available.",
        "derived": "(derived)",
        "version": "Version",
        "locked": "locked",
        "generated_on": "generated",
        "structure_intro": "The following diagram shows how {name} is organized: folders with "
                           "their access scope, and the VIs, type definitions and classes they contain.",
        "uml_intro": "The following class diagram shows the inheritance relationships and the "
                     "methods of each class.",
        "terminals": "Terminals",
        "no_terminals": "No terminals on the connector pane.",
        "t_name": "Name",
        "t_type": "Type",
        "t_dir": "Direction",
        "t_default": "Default",
        "t_desc": "Description",
        "dir_in": "Input",
        "dir_out": "Output",
        "fp_only": "front panel only",
        "req_required": "required",
        "req_recommended": "recommended",
        "req_optional": "optional",
        "non_public": "Non-public members",
        "non_public_intro": "These members are not public and are therefore not documented "
                            "individually.",
        "unreadable": "Unreadable files",
        "missing": "Missing files",
        "notes": "Notes on this run",
        "np_name": "Name",
        "np_scope": "Scope",
        "np_folder": "Folder",
        "u_reason": "Reason",
        "reason_typedef": "type definition (.ctl) — LabVIEW cannot read it as a VI",
        "reason_password": "password-protected",
        "reason_diagram": "block diagram not readable",
        "short_password": "protected",
        "short_diagram": "diagram unreadable",
        "scope_public": "Public",
        "scope_private": "Private",
        "scope_protected": "Protected",
        "scope_community": "Community",
        "scope_unknown": "unknown",
        "legend_folder": "Folder",
        "legend_vi": "VI",
        "legend_vim": "Malleable VI",
        "legend_ctl": "Type definition",
        "legend_class": "Class",
        "legend_inherits": "inherits from",
        "legend_external": "external class",
        "continued": "continued",
        "collapsed": "{n} members — see appendix",
        "structure_collapsed_note": "Non-public folders are shown collapsed; their members are "
                                    "listed individually in the appendix.",
        "more": "{n} more",
        "page": "Page",
        "page_of": "of",
        "items_count": "items",
        "no_images": "Icon and connector pane could not be produced.",
    },
}

REASON_KEY = {"typedef": "reason_typedef", "password": "reason_password",
              "diagram-withheld": "reason_diagram"}
# In the diagram the full sentence would dominate the row; a typedef needs no
# note at all there because its kind color and the legend already say so.
SHORT_REASON = {"password": "short_password", "diagram-withheld": "short_diagram"}

# Rough advance width of Segoe UI, in px per pt of font size. Used to size SVG
# boxes; deliberately generous so text never touches a border.
CHAR_W = 0.54

# Diagrams are laid out in logical px and rasterised at this factor. The SVG's
# width/height carry the factor while the viewBox stays logical, so the drawing
# fills the headless window exactly — emit it at 1x into a 2x window and you get
# half the resolution and three quarters of the canvas blank.
SVG_SCALE = 2


def text_w(s, size):
    return len(s) * size * CHAR_W


def program_files_roots():
    """The Program Files roots to search, newest-API first, no hardcoded drive."""
    roots, seen = [], set()
    for var in ("ProgramW6432", "ProgramFiles", "ProgramFiles(x86)"):
        p = (os.environ.get(var) or "").strip()
        if not p:
            continue
        key = os.path.normcase(p.rstrip("\\/"))
        if key in seen or not os.path.isdir(p):
            continue
        seen.add(key)
        roots.append(p)
    return roots


BROWSER_RELATIVE_PATHS = (
    r"Microsoft\Edge\Application\msedge.exe",
    r"Google\Chrome\Application\chrome.exe",
)


def browser_candidates():
    cands = [os.path.join(root, rel)
             for rel in BROWSER_RELATIVE_PATHS
             for root in program_files_roots()]
    cands += [p for p in (shutil.which("msedge"), shutil.which("chrome")) if p]
    return cands


def render_svg_to_png(svg, W, H, png_path, scale=SVG_SCALE, browser=None):
    candidates = browser_candidates()
    exe = browser or next((p for p in candidates if os.path.isfile(p)), None)
    if not exe:
        raise RuntimeError("No Chromium browser found for diagram rendering. Checked: "
                           + "; ".join(candidates))
    tmp = tempfile.mkdtemp(prefix="lvdoc_")
    try:
        html_path = os.path.join(tmp, "diagram.html")
        with open(html_path, "w", encoding="utf-8") as f:
            f.write("<!doctype html><html><head><meta charset='utf-8'>"
                    "<style>html,body{margin:0;padding:0;background:#fff}svg{display:block}</style>"
                    "</head><body>" + svg + "</body></html>")
        cmd = [
            exe, "--headless=new", "--disable-gpu", "--no-first-run", "--hide-scrollbars",
            "--user-data-dir=" + os.path.join(tmp, "profile"),
            "--screenshot=" + png_path,
            f"--window-size={int(W * scale)},{int(H * scale)}",
            "file:///" + html_path.replace("\\", "/"),
        ]
        subprocess.run(cmd, capture_output=True, timeout=180, check=False)
        if not os.path.isfile(png_path) or os.path.getsize(png_path) == 0:
            raise RuntimeError("Headless browser did not produce the diagram PNG: " + png_path)
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


# --------------------------------------------------------------------------
# Structure diagram: an indented tree, split into balanced columns
# --------------------------------------------------------------------------

RH = 23.0            # row height
INDENT = 20.0        # px per tree level
GUTTER = 34.0
MARGIN = 26.0

PX_PER_CM = 96.0 / 2.54          # 1 CSS px @96 dpi

# Where a diagram may land, as (text width, text height) in cm. Portrait shares
# page 1 with the title block; landscape gets a page of its own, so it reserves
# only the heading and intro. A structure tree of any size is unreadable if it
# is simply squeezed into the portrait text width — choosing the column count
# and the orientation together is what keeps the type large enough to read.
AREA_PORTRAIT = (16.6, 29.7 - 2.2 - 2.0 - 8.0)
AREA_LANDSCAPE = (25.3, 21.0 - 2.2 - 2.0 - 3.2)

# Below this the document stays portrait even when landscape scores higher:
# a rotated page costs the reader something, so it has to buy real legibility.
LANDSCAPE_GAIN = 1.15

# Item rows are set at 11.5 CSS px = 8.6 pt; scaled onto the page they must stay
# above ~6 pt or the tree is a texture, not a diagram. 6 / 8.6 = 0.70. Below
# that the non-public subtrees get collapsed and the plan is recomputed.
ITEM_PT = 11.5 * 0.75
MIN_READABLE_PT = 6.0
MIN_FACTOR = MIN_READABLE_PT / ITEM_PT


def count_leaves(items):
    n = 0
    for it in items or []:
        n += 1
        n += count_leaves(it.get("children"))
    return n


def flatten_structure(data, labels, collapse_non_public=False):
    """The structure tree as a flat row list: (depth, kind, name, scope, note).

    With collapse_non_public, a non-public folder keeps its own row but its
    contents are replaced by a count. Nothing is lost — every one of those
    members is listed by name in the appendix — and the picture stays large
    enough to read, which a 200-row tree squeezed onto one page does not.
    """
    rows = [(0, "library", data["title"], "public", data["target"].get("kind", ""))]
    collapsed = []

    def walk(items, depth):
        for it in items or []:
            kind = (it.get("kind") or "other").lower()
            scope = (it.get("scope") or "unknown").lower()
            note = ""
            if it.get("unreadable") in SHORT_REASON:
                note = labels[SHORT_REASON[it["unreadable"]]]
            container = kind in ("folder", "class", "library")
            if container and collapse_non_public and scope != "public":
                n = count_leaves(it.get("children"))
                if n:
                    collapsed.append((it.get("name", "?"), n))
                    rows.append((depth, kind, it.get("name", "?"), scope,
                                 labels["collapsed"].format(n=n)))
                    continue
            rows.append((depth, kind, it.get("name", "?"), scope, note))
            if container:
                walk(it.get("children"), depth + 1)

    walk(data.get("structure"), 1)
    return rows, collapsed


def split_columns(rows, ncols):
    """Balance the rows over ncols columns.

    A column that starts in the middle of a folder gets a muted continuation
    row naming the folder it belongs to, so no item is ever shown without its
    parent — a silent re-parenting would be a lie about the structure.
    """
    if ncols <= 1:
        return [rows]
    per = math.ceil(len(rows) / ncols)
    cols, i = [], 0
    while i < len(rows):
        chunk = rows[i:i + per]
        if i > 0:
            # Name the enclosing container of the first row in this column.
            depth0 = chunk[0][0]
            parent = next((r for r in reversed(rows[:i]) if r[0] < depth0), None)
            if parent is not None:
                chunk = [(max(0, depth0 - 1), "cont", parent[2], parent[3], "")] + chunk
        cols.append(chunk)
        i += per
    return cols


def column_widths(cols, labels):
    widths = []
    for col in cols:
        w = 0.0
        for depth, kind, name, scope, note in col:
            size = 12.5 if kind in ("library", "folder", "cont") else 11.5
            need = depth * INDENT + 20 + text_w(name, size)
            if kind in ("folder", "cont"):
                need += text_w(labels["scope_" + scope] if kind == "folder"
                               else labels["continued"], 9) + 26
            if note:
                need += text_w(note, 9) + 16
            w = max(w, need)
        widths.append(max(190.0, w + 18))
    return widths


def measure_structure(rows, labels, ncols):
    cols = split_columns(rows, ncols)
    widths = column_widths(cols, labels)
    W = MARGIN * 2 + sum(widths) + GUTTER * (len(cols) - 1)
    H = MARGIN * 2 + max(len(c) * RH for c in cols) + 34   # + legend strip
    return W, H


def choose_placement(measure, ncols_candidates):
    """Pick (orientation, ncols) that renders the diagram largest on the page.

    Returns (orientation, ncols, factor). The factor is what the picture will be
    scaled by in Word — the caller prints it, because "the diagram fits" and
    "the diagram is readable" are not the same statement.
    """
    best = None
    for orient, (aw_cm, ah_cm) in (("portrait", AREA_PORTRAIT),
                                   ("landscape", AREA_LANDSCAPE)):
        aw, ah = aw_cm * PX_PER_CM, ah_cm * PX_PER_CM
        for n in ncols_candidates:
            W, H = measure(n)
            factor = min(aw / W, ah / H, 1.0)
            weighted = factor / (LANDSCAPE_GAIN if orient == "landscape" else 1.0)
            if best is None or weighted > best[0] + 1e-9:
                best = (weighted, orient, n, factor)
    return best[1], best[2], best[3]


def make_structure_svg(rows, labels, ncols):
    cols = split_columns(rows, ncols)
    widths = column_widths(cols, labels)
    W = MARGIN * 2 + sum(widths) + GUTTER * (len(cols) - 1)
    H = MARGIN * 2 + max(len(c) * RH for c in cols) + 34          # + legend strip

    p = [f'<svg xmlns="http://www.w3.org/2000/svg" width="{W * SVG_SCALE:.0f}" '
         f'height="{H * SVG_SCALE:.0f}" viewBox="0 0 {W:.0f} {H:.0f}" '
         f'font-family="{FONT}, Segoe UI, sans-serif">',
         f'<rect width="{W:.0f}" height="{H:.0f}" fill="#ffffff"/>']

    x0 = MARGIN
    kinds_seen = set()
    for col, cw in zip(cols, widths):
        y = MARGIN
        # Tree guides: for every row, a light vertical line at each ancestor level.
        open_at = {}
        for idx, (depth, kind, name, scope, note) in enumerate(col):
            # close guides deeper than this row
            for d in [d for d in open_at if d >= depth]:
                y_from = open_at.pop(d)
                p.append(f'<path d="M {x0 + d * INDENT + 7:.0f} {y_from:.0f} '
                         f'L {x0 + d * INDENT + 7:.0f} {y - RH / 2:.0f}" '
                         f'stroke="#{C_RULE}" stroke-width="1"/>')
            cy = y + RH / 2
            x = x0 + depth * INDENT
            muted = scope != "public"
            ink = C_MUTED if muted else C_INK
            kinds_seen.add(kind)

            if depth > 0:
                p.append(f'<path d="M {x - INDENT + 7:.0f} {cy:.0f} L {x + 2:.0f} {cy:.0f}" '
                         f'stroke="#{C_RULE}" stroke-width="1"/>')

            # Only the root gets the filled chip. A nested library or class is a
            # container like a folder — giving it the root style made it read as
            # a second root.
            if kind == "library" and depth == 0:
                w = text_w(name, 13) + 26
                p.append(f'<rect x="{x:.0f}" y="{y + 2:.0f}" width="{w:.0f}" height="{RH - 4:.0f}" '
                         f'rx="5" fill="#{C_ACCENT}"/>')
                p.append(f'<text x="{x + 13:.0f}" y="{cy + 4.5:.0f}" font-size="13" '
                         f'font-weight="600" fill="#ffffff">{xml_escape(name)}</text>')
            elif kind in ("folder", "cont", "library", "class"):
                nested = kind in ("library", "class")
                lbl = labels["continued"] if kind == "cont" else \
                    (note or labels["scope_" + scope])
                w = text_w(name, 12.5) + text_w(lbl, 9) + 44 + (14 if nested else 0)
                dash = ' stroke-dasharray="4 3"' if kind == "cont" else ""
                p.append(f'<rect x="{x:.0f}" y="{y + 2:.0f}" width="{w:.0f}" height="{RH - 4:.0f}" '
                         f'rx="5" fill="#{C_NODE_FILL}" stroke="#{C_NODE_STROKE}" '
                         f'stroke-width="1.1"{dash}/>')
                p.append(f'<rect x="{x:.0f}" y="{y + 2:.0f}" width="4" height="{RH - 4:.0f}" '
                         f'fill="#{SCOPE_COLOR.get(scope, C_MUTED)}"/>')
                tx = x + 13
                if nested:
                    p.append(f'<rect x="{tx:.0f}" y="{cy - 4:.0f}" width="8" height="8" rx="2" '
                             f'fill="#{KIND_COLOR.get(kind, C_MUTED)}"/>')
                    tx += 14
                p.append(f'<text x="{tx:.0f}" y="{cy + 4:.0f}" font-size="12.5" '
                         f'font-weight="600" fill="#{ink}">{xml_escape(name)}</text>')
                p.append(f'<text x="{x + w - 10:.0f}" y="{cy + 3.5:.0f}" font-size="9" '
                         f'text-anchor="end" fill="#{SCOPE_COLOR.get(scope, C_MUTED)}">'
                         f'{xml_escape(lbl)}</text>')
            else:
                col_k = KIND_COLOR.get(kind, KIND_COLOR["other"])
                p.append(f'<rect x="{x + 4:.0f}" y="{cy - 4:.0f}" width="8" height="8" rx="2" '
                         f'fill="#{col_k}" opacity="{0.45 if muted else 1}"/>')
                it_style = ' font-style="italic"' if muted else ""
                p.append(f'<text x="{x + 18:.0f}" y="{cy + 4:.0f}" font-size="11.5" '
                         f'fill="#{ink}"{it_style}>{xml_escape(name)}</text>')
                if note:
                    nx = x + 18 + text_w(name, 11.5) + 10
                    p.append(f'<text x="{nx:.0f}" y="{cy + 3.5:.0f}" font-size="9" '
                             f'fill="#{C_MUTED}">{xml_escape(note)}</text>')

            if kind in ("library", "folder", "class", "cont"):
                open_at[depth] = cy + RH / 2 - RH / 2 + 4
            y += RH
        for d, y_from in open_at.items():
            p.append(f'<path d="M {x0 + d * INDENT + 7:.0f} {y_from:.0f} '
                     f'L {x0 + d * INDENT + 7:.0f} {y - RH / 2:.0f}" '
                     f'stroke="#{C_RULE}" stroke-width="1"/>')
        x0 += cw + GUTTER

    # Legend — only the kinds that actually occur.
    ly = H - 26
    lx = MARGIN
    entries = [(KIND_COLOR["vi"], labels["legend_vi"], "vi"),
               (KIND_COLOR["vim"], labels["legend_vim"], "vim"),
               (KIND_COLOR["ctl"], labels["legend_ctl"], "ctl"),
               (KIND_COLOR["class"], labels["legend_class"], "class")]
    for color, text, key in entries:
        if key not in kinds_seen:
            continue
        p.append(f'<rect x="{lx:.0f}" y="{ly:.0f}" width="8" height="8" rx="2" fill="#{color}"/>')
        p.append(f'<text x="{lx + 14:.0f}" y="{ly + 8:.0f}" font-size="10.5" '
                 f'fill="#{C_MUTED}">{xml_escape(text)}</text>')
        lx += 14 + text_w(text, 10.5) + 22
    for scope in ("public", "private", "protected", "community"):
        if not any(r[3] == scope for r in rows):
            continue
        p.append(f'<rect x="{lx:.0f}" y="{ly - 1:.0f}" width="4" height="10" '
                 f'fill="#{SCOPE_COLOR[scope]}"/>')
        t = labels["scope_" + scope]
        p.append(f'<text x="{lx + 10:.0f}" y="{ly + 8:.0f}" font-size="10.5" '
                 f'fill="#{C_MUTED}">{xml_escape(t)}</text>')
        lx += 10 + text_w(t, 10.5) + 22

    p.append("</svg>")
    return "".join(p), W, H, len(cols)


# --------------------------------------------------------------------------
# UML class diagram: tidy forest layout, generalization arrows
# --------------------------------------------------------------------------

MAX_METHODS = 14
BOX_GAP_X = 34.0
BOX_GAP_Y = 56.0
LINE_H = 15.0


class UmlBox:
    def __init__(self, name, stereotype, attrs, methods, external, hidden):
        self.name = name
        self.stereotype = stereotype
        self.attrs = attrs
        self.methods = methods          # list of (mark, text, italic)
        self.external = external
        self.hidden = hidden            # methods not shown
        self.children = []
        self.parent = None
        self.x = self.y = 0.0
        title_w = max(text_w(name, 12.5), text_w(stereotype, 9.5) if stereotype else 0)
        body_w = max([text_w(a, 10.5) for a in attrs] +
                     [text_w(m[0] + " " + m[1], 10.5) for m in methods] + [0])
        self.w = max(150.0, title_w + 30, body_w + 30)
        self.head_h = 28.0 + (12.0 if stereotype else 0.0)
        self.attr_h = max(LINE_H, LINE_H * len(attrs)) if attrs else LINE_H * 0.6
        rows = len(methods) + (1 if hidden else 0)
        self.meth_h = max(LINE_H, LINE_H * rows) if rows else LINE_H * 0.6
        self.h = self.head_h + self.attr_h + self.meth_h + 10


def build_uml_boxes(data, labels):
    boxes, truncated = {}, []
    for cls in data.get("classes", []):
        name = cls.get("name", "?")
        methods, shown = [], 0
        for m in cls.get("methods", []):
            if shown >= MAX_METHODS:
                break
            scope = (m.get("scope") or "unknown").lower()
            methods.append((SCOPE_MARK.get(scope, "?"), m.get("name", "?"),
                            bool(m.get("dynamic_dispatch"))))
            shown += 1
        hidden = max(0, len(cls.get("methods", [])) - shown)
        if hidden:
            truncated.append((name, shown, len(cls.get("methods", []))))
        attrs = []
        if cls.get("private_data"):
            attrs.append("- " + cls["private_data"])
        boxes[name] = UmlBox(name, cls.get("stereotype") or "",
                             attrs, methods, False, hidden)

    # External parents get a minimal dashed box so the edge is never dropped.
    for cls in data.get("classes", []):
        parent = (cls.get("parent") or "").strip()
        if not parent:
            continue
        if parent not in boxes:
            boxes[parent] = UmlBox(parent, labels["legend_external"], [], [], True, 0)
        boxes[cls["name"]].parent = boxes[parent]
        boxes[parent].children.append(boxes[cls["name"]])
    return boxes, truncated


def layout_uml(boxes, labels):
    """Tidy forest: leaves get consecutive slots, parents are centered."""
    roots = [b for b in boxes.values() if b.parent is None]

    def depth(b):
        return 0 if b.parent is None else depth(b.parent) + 1

    for b in boxes.values():
        b.level = depth(b)
    level_h = {}
    for b in boxes.values():
        level_h[b.level] = max(level_h.get(b.level, 0.0), b.h)
    y_of, y = {}, MARGIN
    for lvl in sorted(level_h):
        y_of[lvl] = y
        y += level_h[lvl] + BOX_GAP_Y

    cursor = [MARGIN]

    def place(b):
        if not b.children:
            b.x = cursor[0]
            cursor[0] += b.w + BOX_GAP_X
        else:
            for c in b.children:
                place(c)
            first, last = b.children[0], b.children[-1]
            b.x = (first.x + last.x + last.w) / 2 - b.w / 2
            if b.x < cursor[0] - (b.w + BOX_GAP_X):     # keep siblings apart
                b.x = max(b.x, MARGIN)
            cursor[0] = max(cursor[0], b.x + b.w + BOX_GAP_X)
        b.y = y_of[b.level]

    for r in sorted(roots, key=lambda b: b.name.lower()):
        place(r)

    # A parent centred over a narrower subtree can land left of the margin, and
    # the width below is measured from right edges only - so the overhang used to
    # be clipped away. Shift the whole forest back into the canvas instead.
    min_x = min((b.x for b in boxes.values()), default=MARGIN)
    if min_x < MARGIN:
        for b in boxes.values():
            b.x += MARGIN - min_x

    legend_w = MARGIN * 2 + sum(
        text_w(f"{mark} {labels[key]}", 10.5) + 22
        for mark, key in (("+", "scope_public"), ("#", "scope_protected"),
                          ("-", "scope_private"), ("~", "scope_community")))
    W = max(max((b.x + b.w for b in boxes.values()), default=200.0) + MARGIN, legend_w)
    H = max((b.y + b.h for b in boxes.values()), default=200.0) + MARGIN + 26
    return W, H


def make_uml_svg(boxes, labels, W, H):
    p = [f'<svg xmlns="http://www.w3.org/2000/svg" width="{W * SVG_SCALE:.0f}" '
         f'height="{H * SVG_SCALE:.0f}" viewBox="0 0 {W:.0f} {H:.0f}" '
         f'font-family="{FONT}, Segoe UI, sans-serif">',
         f'<rect width="{W:.0f}" height="{H:.0f}" fill="#ffffff"/>',
         '<defs><marker id="gen" viewBox="0 0 12 12" refX="11" refY="6" markerWidth="11" '
         f'markerHeight="11" orient="auto"><path d="M 0 0 L 12 6 L 0 12 z" fill="#ffffff" '
         f'stroke="#{C_EDGE}" stroke-width="1.3"/></marker></defs>']

    # Generalization edges first, so boxes sit on top of them.
    for b in boxes.values():
        if not b.parent:
            continue
        x1, y1 = b.x + b.w / 2, b.y
        x2, y2 = b.parent.x + b.parent.w / 2, b.parent.y + b.parent.h
        ym = (y1 + y2) / 2
        p.append(f'<path d="M {x1:.0f} {y1:.0f} L {x1:.0f} {ym:.0f} L {x2:.0f} {ym:.0f} '
                 f'L {x2:.0f} {y2 + 12:.0f}" fill="none" stroke="#{C_EDGE}" stroke-width="1.4" '
                 'marker-end="url(#gen)"/>')

    for b in boxes.values():
        dash = ' stroke-dasharray="5 3"' if b.external else ""
        p.append(f'<rect x="{b.x:.0f}" y="{b.y:.0f}" width="{b.w:.0f}" height="{b.h:.0f}" rx="4" '
                 f'fill="#ffffff" stroke="#{C_NODE_STROKE}" stroke-width="1.3"{dash}/>')
        p.append(f'<rect x="{b.x:.0f}" y="{b.y:.0f}" width="{b.w:.0f}" height="{b.head_h:.0f}" '
                 f'rx="4" fill="#{C_NODE_FILL}"/>')
        p.append(f'<rect x="{b.x:.0f}" y="{b.y + b.head_h - 4:.0f}" width="{b.w:.0f}" height="4" '
                 f'fill="#{C_NODE_FILL}"/>')
        cx = b.x + b.w / 2
        ty = b.y + 18
        if b.stereotype:
            p.append(f'<text x="{cx:.0f}" y="{ty:.0f}" font-size="9.5" text-anchor="middle" '
                     f'fill="#{C_MUTED}">«{xml_escape(b.stereotype)}»</text>')
            ty += 13
        p.append(f'<text x="{cx:.0f}" y="{ty:.0f}" font-size="12.5" font-weight="600" '
                 f'text-anchor="middle" fill="#{C_MUTED if b.external else C_ACCENT}">'
                 f'{xml_escape(b.name)}</text>')
        p.append(f'<path d="M {b.x:.0f} {b.y + b.head_h:.0f} L {b.x + b.w:.0f} '
                 f'{b.y + b.head_h:.0f}" stroke="#{C_NODE_STROKE}" stroke-width="1"/>')

        y = b.y + b.head_h + 13
        for a in b.attrs:
            p.append(f'<text x="{b.x + 11:.0f}" y="{y:.0f}" font-size="10.5" '
                     f'fill="#{C_MUTED}">{xml_escape(a)}</text>')
            y += LINE_H
        y_sep = b.y + b.head_h + b.attr_h
        p.append(f'<path d="M {b.x:.0f} {y_sep:.0f} L {b.x + b.w:.0f} {y_sep:.0f}" '
                 f'stroke="#{C_NODE_STROKE}" stroke-width="1"/>')
        y = y_sep + 13
        for mark, name, italic in b.methods:
            style = ' font-style="italic"' if italic else ""
            p.append(f'<text x="{b.x + 11:.0f}" y="{y:.0f}" font-size="10.5" '
                     f'fill="#{C_INK}"{style}><tspan fill="#{C_MUTED}">{xml_escape(mark)} </tspan>'
                     f'{xml_escape(name)}</text>')
            y += LINE_H
        if b.hidden:
            p.append(f'<text x="{b.x + 11:.0f}" y="{y:.0f}" font-size="10" '
                     f'fill="#{C_MUTED}">… {xml_escape(labels["more"].format(n=b.hidden))}</text>')

    ly = H - 14
    lx = MARGIN
    for mark, key in (("+", "scope_public"), ("#", "scope_protected"),
                      ("-", "scope_private"), ("~", "scope_community")):
        t = f"{mark} {labels[key]}"
        p.append(f'<text x="{lx:.0f}" y="{ly:.0f}" font-size="10.5" '
                 f'fill="#{C_MUTED}">{xml_escape(t)}</text>')
        lx += text_w(t, 10.5) + 22
    p.append("</svg>")
    return "".join(p)


# --------------------------------------------------------------------------
# DOCX helpers (shared vocabulary with generate_teststand_doc.py)
# --------------------------------------------------------------------------

def _strip_theme(style):
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for att in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        rfonts.attrib.pop(qn("w:" + att), None)
    color = rpr.find(qn("w:color"))
    if color is not None:
        for att in ("themeColor", "themeTint", "themeShade"):
            color.attrib.pop(qn("w:" + att), None)


def tune_style(style, size, color, bold=None, before=None, after=None):
    f = style.font
    f.name = FONT
    f.size = Pt(size)
    f.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        f.bold = bold
    _strip_theme(style)
    pf = getattr(style, "paragraph_format", None)
    if pf is not None:
        if before is not None:
            pf.space_before = Pt(before)
        if after is not None:
            pf.space_after = Pt(after)


def fmt_run(run, size, color, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.bold = bold
    run.font.italic = italic
    return run


def add_rule(doc, color=C_RULE, sz=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    ppr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    ppr.append(pbdr)
    return p


def add_field(paragraph, instruction, placeholder=None, placeholder_fmt=None):
    """Insert a Word field (TOC, PAGE, ...) via fldChar runs."""
    r = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    r._r.append(fld)
    r = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " " + instruction + " "
    r._r.append(instr)
    r = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "separate")
    r._r.append(fld)
    if placeholder:
        r = paragraph.add_run(placeholder)
        if placeholder_fmt:
            placeholder_fmt(r)
    r = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "end")
    r._r.append(fld)


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcpr.append(shd)


def set_table_borders(table):
    """Horizontal lines only — no vertical grid — for a light, modern table."""
    borders = OxmlElement("w:tblBorders")
    for tag, val, sz, color in (
        ("top", "nil", "0", "auto"),
        ("left", "nil", "0", "auto"),
        ("right", "nil", "0", "auto"),
        ("bottom", "single", "4", C_TBL_LINE2),
        ("insideH", "single", "4", C_TBL_LINE),
        ("insideV", "nil", "0", "auto"),
    ):
        el = OxmlElement("w:" + tag)
        el.set(qn("w:val"), val)
        if val != "nil":
            el.set(qn("w:sz"), sz)
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
        borders.append(el)
    table._tbl.tblPr.append(borders)


def set_no_borders(table):
    borders = OxmlElement("w:tblBorders")
    for tag in ("top", "left", "right", "bottom", "insideH", "insideV"):
        el = OxmlElement("w:" + tag)
        el.set(qn("w:val"), "nil")
        borders.append(el)
    table._tbl.tblPr.append(borders)


def set_fixed_layout(table):
    table.allow_autofit = False
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table._tbl.tblPr.append(layout)


def set_cell_margins(table, top=50, bottom=50, left=110, right=110):
    mar = OxmlElement("w:tblCellMar")
    for tag, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement("w:" + tag)
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    table._tbl.tblPr.append(mar)


def mark_header_row(row):
    trpr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trpr.append(el)


def put_cell(cell, text, size=TABLE_PT, color=C_INK, bold=False, italic=False):
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    fmt_run(p.add_run(text), size, color, bold=bold, italic=italic)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def header_table(doc, headers, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    set_fixed_layout(table)
    set_table_borders(table)
    set_cell_margins(table, top=40, bottom=40, left=90, right=90)
    hdr = table.rows[0]
    mark_header_row(hdr)
    hdr.height = Cm(0.56)
    hdr.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    for cell, text, w in zip(hdr.cells, headers, widths):
        cell.width = Cm(w)
        set_cell_shading(cell, C_ACCENT)
        put_cell(cell, text, color="FFFFFF", bold=True)
    return table


def add_body_row(table, values, widths, muted=False):
    row = table.add_row()
    for cell, text, w in zip(row.cells, values, widths):
        cell.width = Cm(w)
        put_cell(cell, text, color=C_MUTED if muted else C_INK)
    return row


def image_size_cm(path, max_w_cm, fallback_w_cm):
    """Natural size in cm at 96 dpi, capped to max_w_cm. Falls back without Pillow."""
    if Image is None or not path or not os.path.isfile(path):
        return Cm(fallback_w_cm)
    try:
        with Image.open(path) as im:
            w_px = im.width or 1
    except Exception:
        return Cm(fallback_w_cm)
    w_cm = w_px / 96.0 * 2.54
    return Cm(min(max(w_cm, 0.6), max_w_cm))


def place_diagram(doc, sec, png_path, w_px, h_px, reserve_cm=2.0):
    """Center the PNG, scaled to fit the text area in both directions."""
    usable_w = sec.page_width - sec.left_margin - sec.right_margin
    usable_h = sec.page_height - sec.top_margin - sec.bottom_margin - Cm(reserve_cm)
    natural_w = Emu(int(w_px * 9525))          # 1 CSS px @96dpi = 9525 EMU
    natural_h = Emu(int(h_px * 9525))
    factor = min(1.0, usable_w / natural_w, usable_h / natural_h)
    doc.add_picture(png_path, width=Emu(int(natural_w * factor)))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_oriented_section(doc, orientation):
    """A new page in the given orientation, inheriting the margins."""
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    if orientation == "landscape":
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width, sec.page_height = Cm(29.7), Cm(21.0)
    else:
        sec.orientation = WD_ORIENT.PORTRAIT
        sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.2)
    sec.top_margin, sec.bottom_margin = Cm(2.2), Cm(2.0)
    return sec


def diagram_chapter(doc, sec, title, intro, png, w_px, h_px, orientation, reserve_cm):
    """A diagram with its heading, rotated onto its own page when that helps.

    The heading is a styled paragraph rather than a Heading style so the chapter
    stays out of the table of contents, which then lists only the per-VI sections.
    """
    if orientation == "landscape":
        sec = add_oriented_section(doc, "landscape")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0 if orientation == "landscape" else 16)
    p.paragraph_format.space_after = Pt(6)
    fmt_run(p.add_run(title), 15, C_ACCENT, bold=True)
    fmt_run(doc.add_paragraph().add_run(intro), 10.5, C_INK)
    place_diagram(doc, sec, png, w_px, h_px, reserve_cm=reserve_cm)
    if orientation == "landscape":
        sec = add_oriented_section(doc, "portrait")
    return sec


# --------------------------------------------------------------------------
# Document assembly
# --------------------------------------------------------------------------

def build_document(data, labels, diagrams, out_path):
    doc = Document()

    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.2)
    sec.top_margin, sec.bottom_margin = Cm(2.2), Cm(2.0)

    tune_style(doc.styles["Normal"], 10.5, C_INK, after=6)
    tune_style(doc.styles["Heading 1"], 15, C_ACCENT, bold=True, before=18, after=8)
    tune_style(doc.styles["Heading 2"], 12.5, C_ACCENT2, bold=True, before=14, after=5)
    tune_style(doc.styles["Heading 3"], 11, C_ACCENT2, bold=True, before=10, after=4)

    target = data.get("target", {})
    vis = data.get("vis", [])

    # --- Header block --------------------------------------------------------
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    fmt_run(p.add_run(data["title"]), 26, C_ACCENT, bold=True)

    meta = [target.get("path", "")]
    if target.get("version"):
        meta.append(f"{labels['version']} {target['version']}")
    if target.get("labview"):
        meta.append("LabVIEW " + str(target["labview"]))
    meta.append(f"{len(vis)} {labels['public_vis']}")
    if target.get("locked"):
        meta.append(labels["locked"])
    if data.get("generated"):
        meta.append(f"{labels['generated_on']} {data['generated']}")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    fmt_run(p.add_run("  ·  ".join(x for x in meta if x)), 9, C_MUTED)
    add_rule(doc)

    desc = (target.get("description") or "").strip()
    p = doc.add_paragraph()
    if desc:
        fmt_run(p.add_run(desc), 10.5, C_INK)
    else:
        fmt_run(p.add_run(labels["no_description"]), 10.5, C_MUTED, italic=True)

    # --- Structure (up front, kept out of the TOC like the title block) ------
    # A rotated chapter ends on a fresh page by itself; adding a page break on
    # top of that section break would leave a blank page behind.
    s_land = diagrams["structure_orientation"] == "landscape"
    intro = labels["structure_intro"].format(name=data["title"])
    if data.get("structure_note"):
        intro += " " + data["structure_note"]
    sec = diagram_chapter(
        doc, sec, labels["structure"], intro,
        diagrams["structure"], diagrams["structure_w"], diagrams["structure_h"],
        diagrams["structure_orientation"], reserve_cm=3.2 if s_land else 8.0)
    on_fresh_page = s_land

    # --- UML -----------------------------------------------------------------
    if diagrams.get("uml"):
        u_land = diagrams["uml_orientation"] == "landscape"
        if not u_land and not on_fresh_page:
            doc.add_page_break()
        sec = diagram_chapter(
            doc, sec, labels["uml"], labels["uml_intro"],
            diagrams["uml"], diagrams["uml_w"], diagrams["uml_h"],
            diagrams["uml_orientation"], reserve_cm=3.2)
        on_fresh_page = u_land

    # --- Table of contents ---------------------------------------------------
    if not on_fresh_page:
        doc.add_page_break()
    p = doc.add_paragraph()  # deliberately NOT a Heading style: keeps it out of the TOC
    p.paragraph_format.space_after = Pt(10)
    fmt_run(p.add_run(labels["toc"]), 15, C_ACCENT, bold=True)
    add_field(
        doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
        placeholder=labels["toc_placeholder"],
        placeholder_fmt=lambda r: fmt_run(r, 9.5, C_MUTED, italic=True),
    )
    doc.add_page_break()

    # --- One section per public VI -------------------------------------------
    doc.add_heading(labels["public_vis"], level=1)
    # No connector-pane index column: the terminals are already ordered by index,
    # the picture above the table is labelled with them, and a terminal that is
    # NOT on the pane says so in the direction column.
    t_widths = (3.6, 2.6, 3.0, 2.4, 5.0)
    t_headers = (labels["t_name"], labels["t_type"], labels["t_dir"],
                 labels["t_default"], labels["t_desc"])

    for vi in vis:
        doc.add_heading(vi.get("name", "?"), level=2)

        icon, conpane = vi.get("icon"), vi.get("conpane")
        icon_ok = bool(icon) and os.path.isfile(icon)
        cp_ok = bool(conpane) and os.path.isfile(conpane)
        if icon_ok or cp_ok:
            same = icon_ok and cp_ok and os.path.normcase(icon) == os.path.normcase(conpane)
            cells_needed = 1 if same else (int(icon_ok) + int(cp_ok))
            tbl = doc.add_table(rows=1, cols=max(1, cells_needed))
            tbl.autofit = False
            set_no_borders(tbl)
            set_cell_margins(tbl, top=0, bottom=60, left=0, right=200)
            cells = tbl.rows[0].cells
            slot = 0
            if icon_ok:
                cells[slot].width = Cm(2.4)
                run = cells[slot].paragraphs[0].add_run()
                run.add_picture(icon, width=image_size_cm(icon, 2.2, 1.3))
                slot += 1
            if cp_ok and not same:
                cells[slot].width = Cm(13.0)
                run = cells[slot].paragraphs[0].add_run()
                run.add_picture(conpane, width=image_size_cm(conpane, 8.0, 5.0))

        d = (vi.get("description") or "").strip()
        p = doc.add_paragraph()
        if d:
            fmt_run(p.add_run(d), 10.5, C_INK)
            if vi.get("description_derived"):
                fmt_run(p.add_run("  " + labels["derived"]), 8.5, C_MUTED, italic=True)
        else:
            fmt_run(p.add_run(labels["no_description"]), 10.5, C_MUTED, italic=True)

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        fmt_run(p.add_run(labels["terminals"]), 10, C_ACCENT, bold=True)

        terms = list(vi.get("terminals") or [])
        # Connector-pane terminals first, ordered by index; front-panel-only last.
        terms.sort(key=lambda t: (t.get("conIdx") is None, t.get("conIdx") if
                                  isinstance(t.get("conIdx"), int) else 0))
        if not terms:
            fmt_run(doc.add_paragraph().add_run(labels["no_terminals"]), 9.5, C_MUTED, italic=True)
        else:
            table = header_table(doc, t_headers, t_widths)
            for t in terms:
                idx = t.get("conIdx")
                on_pane = isinstance(idx, int)
                direction = labels["dir_in"] if (t.get("direction") or "").lower() == "input" \
                    else labels["dir_out"]
                # AIXML's connection attribute: required / recommended / optional.
                # Without a connector-pane picture this is the only place the
                # reader learns which terminals must be wired.
                req = labels.get("req_" + (t.get("required") or "").lower())
                if req:
                    direction = f"{direction} ({req})"
                add_body_row(table, (
                    t.get("name", ""),
                    t.get("type", "") or "—",
                    direction if on_pane else labels["fp_only"],
                    str(t.get("default", "")).strip() or "—",
                    (t.get("description") or "").strip() or "—",
                ), t_widths, muted=not on_pane)
            doc.add_paragraph().paragraph_format.space_after = Pt(0)

    # --- Appendix -------------------------------------------------------------
    non_public = data.get("non_public") or []
    unreadable = data.get("unreadable") or []
    missing = data.get("missing_files") or []
    notes = data.get("notes") or []
    if non_public or unreadable or missing or notes:
        doc.add_page_break()
        doc.add_heading(labels["appendix"], level=1)

    if non_public:
        doc.add_heading(labels["non_public"], level=2)
        fmt_run(doc.add_paragraph().add_run(labels["non_public_intro"]), 10.5, C_INK)
        w = (8.0, 3.4, 5.2)
        table = header_table(doc, (labels["np_name"], labels["np_scope"], labels["np_folder"]), w)
        for it in non_public:
            scope = (it.get("scope") or "unknown").lower()
            add_body_row(table, (it.get("name", ""), labels.get("scope_" + scope, scope),
                                 it.get("folder", "") or "—"), w)
        doc.add_paragraph().paragraph_format.space_after = Pt(0)

    if unreadable:
        doc.add_heading(labels["unreadable"], level=2)
        w = (7.0, 9.6)
        table = header_table(doc, (labels["np_name"], labels["u_reason"]), w)
        for it in unreadable:
            reason = it.get("reason", "")
            add_body_row(table, (it.get("name", ""),
                                 labels.get(REASON_KEY.get(reason, ""), reason)), w)
        doc.add_paragraph().paragraph_format.space_after = Pt(0)

    if missing:
        doc.add_heading(labels["missing"], level=2)
        for m in missing:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            fmt_run(p.add_run(m), 9.5, C_MUTED)

    if notes:
        doc.add_heading(labels["notes"], level=2)
        for n in notes:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            fmt_run(p.add_run("· " + n), 9.5, C_MUTED)

    # --- Footer ---------------------------------------------------------------
    footer_p = sec.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fmt_run(footer_p.add_run(labels["page"] + " "), 8.5, C_MUTED)
    add_field(footer_p, "PAGE")
    fmt_run(footer_p.add_run(" " + labels["page_of"] + " "), 8.5, C_MUTED)
    add_field(footer_p, "NUMPAGES")
    for run in footer_p.runs:
        fmt_run(run, 8.5, C_MUTED, bold=run.font.bold or False)

    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    doc.settings.element.append(upd)

    doc.save(out_path)


# --------------------------------------------------------------------------
# Main
# --------------------------------------------------------------------------

def main(argv):
    args = [a for a in argv if not a.startswith("--")]
    if len(args) < 2:
        print(__doc__)
        return 1
    data_path, out_path = os.path.abspath(args[0]), os.path.abspath(args[1])

    def opt(name):
        return argv[argv.index(name) + 1] if name in argv else None

    with open(data_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    labels = dict(LABELS.get(str(data.get("language", DEFAULT_LANGUAGE)).lower(),
                             LABELS[DEFAULT_LANGUAGE]))
    labels.update(data.get("labels", {}))
    browser = opt("--browser")

    diagrams = {}

    rows, collapsed = flatten_structure(data, labels)
    s_orient, ncols, s_factor = choose_placement(
        lambda n: measure_structure(rows, labels, n), range(1, 6))
    if s_factor < MIN_FACTOR:
        # Too small to read. Collapse the non-public subtrees (their members are
        # in the appendix either way) and try again; keep whichever reads better.
        rows2, collapsed2 = flatten_structure(data, labels, collapse_non_public=True)
        orient2, ncols2, factor2 = choose_placement(
            lambda n: measure_structure(rows2, labels, n), range(1, 6))
        if factor2 > s_factor:
            rows, collapsed = rows2, collapsed2
            s_orient, ncols, s_factor = orient2, ncols2, factor2
    if collapsed:
        data = dict(data, structure_note=labels["structure_collapsed_note"])
    svg, sw, sh, ncols = make_structure_svg(rows, labels, ncols)
    structure_png = os.path.abspath(opt("--structure-out") or os.path.join(
        tempfile.gettempdir(), "lvdoc_structure.png"))
    render_svg_to_png(svg, sw, sh, structure_png, browser=browser)
    diagrams.update(structure=structure_png, structure_w=sw, structure_h=sh,
                    structure_orientation=s_orient)

    truncated = []
    if data.get("classes"):
        boxes, truncated = build_uml_boxes(data, labels)
        uw, uh = layout_uml(boxes, labels)
        u_orient, _, u_factor = choose_placement(lambda _n: (uw, uh), (1,))
        usvg = make_uml_svg(boxes, labels, uw, uh)
        uml_png = os.path.abspath(opt("--uml-out") or os.path.join(
            tempfile.gettempdir(), "lvdoc_uml.png"))
        render_svg_to_png(usvg, uw, uh, uml_png, browser=browser)
        diagrams.update(uml=uml_png, uml_w=uw, uml_h=uh, uml_orientation=u_orient)
        n_edges = sum(1 for b in boxes.values() if b.parent)
        n_boxes = len(boxes)

    build_document(data, labels, diagrams, out_path)

    vis = data.get("vis", [])
    n_icons = sum(1 for v in vis if v.get("icon") and os.path.isfile(v["icon"]))
    n_panes = sum(1 for v in vis if v.get("conpane") and os.path.isfile(v["conpane"]))
    n_derived = sum(1 for v in vis if v.get("description_derived"))

    print(f"[ok] docx      : {out_path}")
    print(f"[ok] structure : {structure_png} ({int(sw * SVG_SCALE)}x{int(sh * SVG_SCALE)} px, "
          f"{len(rows)} rows, {ncols} column(s), {s_orient}, "
          f"shown at {s_factor * 100:.0f}% = {ITEM_PT * s_factor:.1f} pt"
          f"{'' if s_factor >= MIN_FACTOR else ' — BELOW THE READABLE FLOOR'})")
    for name, n in collapsed:
        print(f"[..] collapsed : {name} ({n} members) shown as a count, listed in the appendix")
    if diagrams.get("uml"):
        print(f"[ok] uml       : {diagrams['uml']} ({int(diagrams['uml_w'] * SVG_SCALE)}x"
              f"{int(diagrams['uml_h'] * SVG_SCALE)} px, {n_boxes} classes, {n_edges} edges, "
              f"{u_orient}, shown at {u_factor * 100:.0f}%)")
    else:
        print("[--] uml       : omitted (no classes in the data)")
    if n_icons or n_panes:
        print(f"[ok] images    : {n_icons} icons, {n_panes} connector panes")
    else:
        print("[--] images    : none supplied (no icon/connector-pane files)")
    print(f"[ok] content   : {len(vis)} public VIs ({n_derived} derived descriptions), "
          f"{len(data.get('non_public') or [])} non-public, "
          f"{len(data.get('classes') or [])} classes, "
          f"{len(data.get('unreadable') or [])} unreadable, "
          f"{len(data.get('missing_files') or [])} missing")
    for name, shown, total in truncated:
        print(f"[..] truncated : {name} shows {shown} of {total} methods in the UML diagram")
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))
